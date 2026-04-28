"""
Generates Rechnung 07/2026 – Teilrechnung TimeToJump Dolomites (V3).
Based on user's manually edited version. Two-page invoice with:
- Pos 1-7: Angebot 02 Leistungen (Status: Fertig)
- Pos 8-18: Besprochene Erweiterungen (Status: In Bearbeitung)
- Offene Betraege aufgeschluesselt nach Angebot 02 + telef. Absprache
"""
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# Load user-edited template to preserve their manual changes (UID etc.)
TEMPLATE = r"G:\Andere Computer\Mein Laptop\Documents\01_Finanzen_und_Vermoegen\Buchhaltung\Buchhaltung Solum Invest GmbH\Buchhaltung GmbH 2026\Dauerrechnung Hoettinger Au 43 Top 8 - ab 05-2026.docx"
doc = Document(TEMPLATE)

# Clear template content
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
for t in doc.tables:
    t._element.getparent().remove(t._element)

# ── Colours ────────────────────────────────────────────────────────
C_TITLE = RGBColor(0x1A, 0x1A, 0x1A)
C_ADDR  = RGBColor(0x99, 0x99, 0x99)
C_SUB   = RGBColor(0x88, 0x88, 0x88)
C_DESC  = RGBColor(0x66, 0x66, 0x66)
C_SECT  = RGBColor(0x55, 0x55, 0x55)
C_HEAD  = RGBColor(0x44, 0x44, 0x44)
C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_GREEN = RGBColor(0x2E, 0x7D, 0x32)
C_AMBER = RGBColor(0xE6, 0x8A, 0x00)

# ── Helpers ────────────────────────────────────────────────────────
def add_p(text, bold=False, size_emu=None, color=None, align=None,
          space_after=None, space_before=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size_emu: run.font.size = Emu(size_emu)
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Emu(space_after)
    if space_before is not None: p.paragraph_format.space_before = Emu(space_before)
    return p, run

def add_run(p, text, bold=False, size_emu=None, color=None, italic=False):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size_emu: run.font.size = Emu(size_emu)
    if color: run.font.color.rgb = color
    return run

def shade_cells(row, color_hex):
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

def make_borderless(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_cell(cell, text, bold=False, size_emu=None, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if size_emu: run.font.size = Emu(size_emu)
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Emu(25400)
    p.paragraph_format.space_before = Emu(12700)

def set_cell_multi(cell, lines, size_emu=101600, color=C_SECT, bold_first=False):
    """Set cell with multiple lines, first line optionally bold."""
    cell.text = ''
    p = cell.paragraphs[0]
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run('\n').font.size = Emu(size_emu)
        run = p.add_run(line)
        run.font.size = Emu(size_emu)
        run.font.color.rgb = color
        if bold_first and i == 0:
            run.bold = True
    p.paragraph_format.space_after = Emu(25400)
    p.paragraph_format.space_before = Emu(12700)

# ══════════════════════════════════════════════════════════════════
#  PAGE 1
# ══════════════════════════════════════════════════════════════════

# ── Company header ─────────────────────────────────────────────────
add_p('Solum Invest GmbH', bold=True, size_emu=228600, color=C_TITLE, space_after=25400)
add_p('Augasse 4, 6111 Volders', size_emu=101600, color=C_ADDR, space_after=152400)
add_p('', space_after=76200)

# ── Header table ───────────────────────────────────────────────────
ht = doc.add_table(rows=5, cols=2)
make_borderless(ht)
ht.columns[0].width = Cm(9)
ht.columns[1].width = Cm(7)

set_cell(ht.rows[0].cells[0], 'An:', size_emu=101600, color=C_DESC)
set_cell(ht.rows[0].cells[1], 'Rechnung Nr. 07/2026', bold=True, size_emu=133350, color=C_TITLE)
set_cell(ht.rows[1].cells[0], 'Time to Jump Dolomites', bold=True, size_emu=114300)
set_cell(ht.rows[1].cells[1], 'Datum: 28. April 2026', size_emu=101600, color=C_DESC)
set_cell(ht.rows[2].cells[0], 'David Prato', size_emu=114300, color=C_SECT)
set_cell(ht.rows[2].cells[1], 'Zahlungsstatus: Bezahlt', bold=True, size_emu=101600, color=C_GREEN)
set_cell(ht.rows[3].cells[0], 'info@timetojumpdolomites.com', size_emu=101600, color=C_DESC)
set_cell(ht.rows[3].cells[1], '', size_emu=101600)
set_cell(ht.rows[4].cells[0], 'UID: IT03113270213', size_emu=101600, color=C_DESC)
set_cell(ht.rows[4].cells[1], '', size_emu=101600)

# ── Title block ────────────────────────────────────────────────────
add_p('RECHNUNG', bold=True, size_emu=254000, color=C_TITLE,
      align=WD_ALIGN_PARAGRAPH.LEFT, space_after=25400, space_before=177800)

add_p('Teilrechnung 1  \u00b7  Projekt: Website Relaunch TimeToJump Dolomites',
      size_emu=101600, color=C_SUB, space_after=76200)

add_p('Referenz: Angebot Nr. 02/2026 vom 15. M\u00e4rz 2026  \u00b7  Gesamtauftragswert: EUR 2.000,\u2013',
      size_emu=101600, color=C_SUB, space_after=101600)

# ── Intro ──────────────────────────────────────────────────────────
p, _ = add_p('', space_after=76200)
add_run(p, 'Sehr geehrter Herr Prato,', bold=True, size_emu=114300)

p, _ = add_p('', space_after=76200)
add_run(p, 'hiermit stellen wir Ihnen die erste Teilrechnung f\u00fcr die im Rahmen des Angebots Nr. 02/2026 sowie der telefonischen Vereinbarung vom 27.04.2026 erbrachten und beauftragten Leistungen.',
        size_emu=114300, color=C_SECT)

# ── Leistungsuebersicht ───────────────────────────────────────────
add_p('Leistungs\u00fcbersicht', bold=True, size_emu=152400, color=C_BLACK,
      space_before=101600, space_after=50800)

# ── Section A: Angebot 02 ─────────────────────────────────────────
add_p('Angebot Nr. 02/2026 \u2013 Website Relaunch', bold=True, size_emu=114300, color=C_HEAD,
      space_after=38100, space_before=50800)

t1 = doc.add_table(rows=1, cols=3)
make_borderless(t1)
t1.columns[0].width = Cm(1.2)
t1.columns[1].width = Cm(11.5)
t1.columns[2].width = Cm(3.3)

set_cell(t1.rows[0].cells[0], 'Pos.', bold=True, size_emu=88900, color=C_HEAD)
set_cell(t1.rows[0].cells[1], 'Beschreibung', bold=True, size_emu=88900, color=C_HEAD)
set_cell(t1.rows[0].cells[2], 'Status', bold=True, size_emu=88900, color=C_HEAD, align=WD_ALIGN_PARAGRAPH.RIGHT)
shade_cells(t1.rows[0], 'F2F2F2')

items_a = [
    ('1', ['Vollst\u00e4ndige Landingpage \u2013 12 Sektionen, responsiv',
           'Hero, About, What to Expect, Pricing, Testimonials, FAQ,',
           'Anforderungen, Google Maps, Footer mit Social-Media-Links']),
    ('2', ['Mehrsprachigkeit \u2013 9 Sprachen mit LocalStorage',
           'EN, DE, IT, LAD, NL, FR, PL, ES, CS']),
    ('3', ['Kontaktformular mit E-Mail-Versand',
           'FormSubmit.co, CAPTCHA + Honeypot, Fehlerbehandlung']),
    ('4', ['WhatsApp / E-Mail / Telefon \u2013 schwebendes Chat-Widget',
           'Direktlinks, sichtbar auf allen Endger\u00e4ten']),
    ('5', ['Sicherheit & SEO-Grundstruktur',
           'Security Headers (CSP), SSL/HTTPS, Meta-Tags, Open Graph,',
           'semantisches HTML5, Ladezeit < 1s']),
    ('6', ['Datenschutzseite (privacy.html)']),
    ('7', ['Feedback-Schleifen vor Go-Live (> 2 absolviert)']),
]

for pos, lines in items_a:
    row = t1.add_row()
    set_cell(row.cells[0], pos, bold=True, size_emu=101600, color=C_SECT)
    set_cell_multi(row.cells[1], lines, size_emu=88900, color=C_SECT, bold_first=True)
    set_cell(row.cells[2], 'Fertig', size_emu=88900, color=C_GREEN, align=WD_ALIGN_PARAGRAPH.RIGHT)

# ── Section B: Telefonische Absprache ──────────────────────────────
add_p('Telefonische Absprache vom 27.04.2026 \u2013 Erweiterungen', bold=True, size_emu=114300, color=C_HEAD,
      space_after=38100, space_before=101600)

t2 = doc.add_table(rows=1, cols=3)
make_borderless(t2)
t2.columns[0].width = Cm(1.2)
t2.columns[1].width = Cm(11.5)
t2.columns[2].width = Cm(3.3)

set_cell(t2.rows[0].cells[0], 'Pos.', bold=True, size_emu=88900, color=C_HEAD)
set_cell(t2.rows[0].cells[1], 'Beschreibung', bold=True, size_emu=88900, color=C_HEAD)
set_cell(t2.rows[0].cells[2], 'Status', bold=True, size_emu=88900, color=C_HEAD, align=WD_ALIGN_PARAGRAPH.RIGHT)
shade_cells(t2.rows[0], 'F2F2F2')

items_b = [
    ('8',  ['Supabase-Backend mit PostgreSQL-Datenbank',
            'Event- & Buchungstabellen, Row Level Security, Edge Functions']),
    ('9',  ['Booking Engine & dynamische Flight Cards',
            'Live-Verf\u00fcgbarkeit, sequenzielle Load-Logik, Nickname-Pills']),
    ('10', ['Dynamisches Gruppenpricing',
            'Preis pro Person sinkt mit jedem weiteren Springer']),
    ('11', ['Buchungsformular (Booking Modal)',
            'Alle Pflichtfelder, Glassmorphism-Design, 9 Sprachen']),
    ('12', ['Stripe-Integration (Kreditkarten-Hold)',
            'SetupIntent, 0\u20ac Card Hold, Strafeinzug bei Stornierung']),
    ('13', ['AGB & Rechtstexte',
            'Stornobedingungen, Briefing-Strafe, Zero Tolerance, Short-Notice']),
    ('14', ['Admin-Dashboard (admin.html)',
            'Event-Verwaltung & Buchungs\u00fcbersicht']),
    ('15', ['Website-Optimierungen & Content-Updates',
            'Pricing, FAQ, Maps, Navigation, Elikos-Bereinigung']),
    ('16', ['Erfolgsseite (success.html)',
            'WhatsApp-Gruppeneinladung nach erfolgreicher Buchung']),
    ('17', ['DNS-Migration & Wix-Abl\u00f6se',
            'Domain-Umzug auf Cloudflare Pages, SSL, Weiterleitungen']),
    ('18', ['Automatische Reminder-E-Mails',
            '7 / 5 / 3 Tage vor Abflug, Gruppenbuchungsfunktion']),
]

for pos, lines in items_b:
    row = t2.add_row()
    set_cell(row.cells[0], pos, bold=True, size_emu=101600, color=C_SECT)
    set_cell_multi(row.cells[1], lines, size_emu=88900, color=C_SECT, bold_first=True)
    set_cell(row.cells[2], 'In Bearbeitung', size_emu=88900, color=C_AMBER, align=WD_ALIGN_PARAGRAPH.RIGHT)

# ══════════════════════════════════════════════════════════════════
#  PREISTABELLE
# ══════════════════════════════════════════════════════════════════
add_p('Preis\u00fcbersicht', bold=True, size_emu=152400, color=C_BLACK,
      space_before=127000, space_after=50800)

pt = doc.add_table(rows=4, cols=2)
make_borderless(pt)
pt.columns[0].width = Cm(12)
pt.columns[1].width = Cm(4)

set_cell(pt.rows[0].cells[0], 'Teilbetrag Positionen 1\u20137 (Angebot Nr. 02/2026)', size_emu=114300, color=C_SECT)
set_cell(pt.rows[0].cells[1], 'EUR 600,\u2013', bold=True, size_emu=114300, align=WD_ALIGN_PARAGRAPH.RIGHT)

set_cell(pt.rows[1].cells[0], 'Umsatzsteuer (Reverse Charge \u2013 Art. 196 MwSt-RL)', size_emu=101600, color=C_DESC)
set_cell(pt.rows[1].cells[1], '0,00 EUR', size_emu=101600, color=C_DESC, align=WD_ALIGN_PARAGRAPH.RIGHT)

set_cell(pt.rows[2].cells[0], '', size_emu=50800)
set_cell(pt.rows[2].cells[1], '', size_emu=50800)

set_cell(pt.rows[3].cells[0], 'Rechnungsbetrag', bold=True, size_emu=133350, color=C_TITLE)
set_cell(pt.rows[3].cells[1], 'EUR 600,\u2013', bold=True, size_emu=133350, color=C_TITLE, align=WD_ALIGN_PARAGRAPH.RIGHT)
shade_cells(pt.rows[3], 'F2F2F2')

# ══════════════════════════════════════════════════════════════════
#  ZAHLUNGSSTATUS
# ══════════════════════════════════════════════════════════════════
add_p('Zahlungsstatus', bold=True, size_emu=127000, color=C_HEAD,
      space_before=152400, space_after=50800)

p, _ = add_p('', space_after=50800)
add_run(p, 'Der Rechnungsbetrag von EUR 600,\u2013 wurde bereits vollst\u00e4ndig beglichen.',
        bold=True, size_emu=114300, color=C_GREEN)

# ── Offene Betraege ────────────────────────────────────────────────
add_p('Offene Betr\u00e4ge', bold=True, size_emu=127000, color=C_HEAD,
      space_before=101600, space_after=50800)

ot = doc.add_table(rows=3, cols=3)
make_borderless(ot)
ot.columns[0].width = Cm(8.5)
ot.columns[1].width = Cm(4)
ot.columns[2].width = Cm(3.5)

set_cell(ot.rows[0].cells[0], 'Verbleibender offener Betrag Angebot Nr. 02/2026',
         size_emu=101600, color=C_SECT)
set_cell(ot.rows[0].cells[1], 'EUR 300,\u2013', bold=True, size_emu=101600, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell(ot.rows[0].cells[2], 'F\u00e4llig Ende Mai', size_emu=88900, color=C_AMBER, align=WD_ALIGN_PARAGRAPH.RIGHT)

set_cell(ot.rows[1].cells[0], 'Verbleibender offener Betrag aus\ntelefonischer Absprache 27.04.2026',
         size_emu=101600, color=C_SECT)
set_cell(ot.rows[1].cells[1], 'EUR 1.100,\u2013', bold=True, size_emu=101600, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell(ot.rows[1].cells[2], 'F\u00e4llig nach Go-Live\nmit DNS-Umstellung', size_emu=88900, color=C_AMBER, align=WD_ALIGN_PARAGRAPH.RIGHT)

shade_cells(ot.rows[2], 'F2F2F2')
set_cell(ot.rows[2].cells[0], 'Gesamt offen', bold=True, size_emu=114300, color=C_TITLE)
set_cell(ot.rows[2].cells[1], 'EUR 1.400,\u2013', bold=True, size_emu=114300, color=C_TITLE, align=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell(ot.rows[2].cells[2], '', size_emu=88900)

# ══════════════════════════════════════════════════════════════════
#  STEUERHINWEIS
# ══════════════════════════════════════════════════════════════════
add_p('Steuerhinweis', bold=True, size_emu=114300, color=C_HEAD,
      space_before=127000, space_after=38100)

add_p('Diese Leistung ist eine innergemeinschaftliche sonstige Leistung (B2B) gem\u00e4\u00df '
      'Art. 196 MwSt-Richtlinie. Es wird keine \u00f6sterreichische Umsatzsteuer in Rechnung '
      'gestellt. Der Leistungsempf\u00e4nger schuldet die Steuer im eigenen Land (Reverse Charge).',
      italic=True, size_emu=101600, color=C_DESC, space_after=101600)

# ══════════════════════════════════════════════════════════════════
#  SIGNATUR & BANK
# ══════════════════════════════════════════════════════════════════
add_p('Mit freundlichen Gr\u00fc\u00dfen,', size_emu=114300, color=C_SECT,
      space_before=127000, space_after=203200)

add_p('Axel Ludwig M\u00fchlbacher', bold=True, size_emu=127000, color=C_TITLE, space_after=25400)
add_p('Gesch\u00e4ftsf\u00fchrer \u2013 Solum Invest GmbH', size_emu=101600, color=C_DESC, space_after=101600)

bt = doc.add_table(rows=4, cols=2)
make_borderless(bt)
bt.columns[0].width = Cm(4)
bt.columns[1].width = Cm(12)

set_cell(bt.rows[0].cells[0], 'Bank:', size_emu=101600, color=C_DESC)
set_cell(bt.rows[0].cells[1], 'Raiffeisenbank Volders', size_emu=101600, color=C_SECT)
set_cell(bt.rows[1].cells[0], 'IBAN:', size_emu=101600, color=C_DESC)
set_cell(bt.rows[1].cells[1], 'AT08 3620 0000 0035 2971', size_emu=101600, color=C_SECT)
set_cell(bt.rows[2].cells[0], 'BIC:', size_emu=101600, color=C_DESC)
set_cell(bt.rows[2].cells[1], 'RZTIAT 22200', size_emu=101600, color=C_SECT)
set_cell(bt.rows[3].cells[0], 'Verwendung:', size_emu=101600, color=C_DESC)
set_cell(bt.rows[3].cells[1], 'Rechnung 07/2026 \u2013 TimeToJump Dolomites', size_emu=101600, color=C_SECT)

# ── Footer ─────────────────────────────────────────────────────────
add_p('Solum Invest GmbH  \u00b7  Augasse 4, 6111 Volders  \u00b7  ATU81665404  \u00b7  FN 642566a  \u00b7  office@soluminvest.at  \u00b7  +43 676 4008771',
      size_emu=88900, color=C_ADDR, align=WD_ALIGN_PARAGRAPH.CENTER,
      space_before=203200, space_after=0)

# ══════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════
FILENAME = 'Rechnung_07_2026_Teilrechnung_TimeToJump_Dolomites.docx'

out1 = os.path.join(r'c:\AlpineDigital\timetojumpdolomites\Korrespondenz\Angebote und Rechnungen', FILENAME)
doc.save(out1)
print(f'Saved: {out1}')

acct_dir = r'G:\Andere Computer\Mein Laptop\Documents\01_Finanzen_und_Vermoegen\Buchhaltung\Buchhaltung Solum Invest GmbH\Buchhaltung GmbH 2026\04-April-2026\Rechnungen Out - April 2026'
os.makedirs(acct_dir, exist_ok=True)
out2 = os.path.join(acct_dir, FILENAME)
doc.save(out2)
print(f'Saved: {out2}')
