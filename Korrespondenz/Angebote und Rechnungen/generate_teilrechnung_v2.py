"""
Generates Teilrechnung 07/2026 for TimeToJump Dolomites project.
Uses the exact same styling as the Dauerrechnung Höttinger Au 43 Top 8 template.
"""
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Load template to clone its base style ──────────────────────────
TEMPLATE = r"G:\Andere Computer\Mein Laptop\Documents\01_Finanzen_und_Vermoegen\Buchhaltung\Buchhaltung Solum Invest GmbH\Buchhaltung GmbH 2026\Dauerrechnung Hoettinger Au 43 Top 8 - ab 05-2026.docx"
doc = Document(TEMPLATE)

# Clear all content from template
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)
for t in doc.tables:
    t._element.getparent().remove(t._element)

# ── Colours (from template) ────────────────────────────────────────
C_TITLE   = RGBColor(0x1A, 0x1A, 0x1A)
C_ADDR    = RGBColor(0x99, 0x99, 0x99)
C_SUB     = RGBColor(0x88, 0x88, 0x88)
C_DESC    = RGBColor(0x66, 0x66, 0x66)
C_SECT    = RGBColor(0x55, 0x55, 0x55)
C_HEAD    = RGBColor(0x44, 0x44, 0x44)
C_BLACK   = RGBColor(0x00, 0x00, 0x00)

# ── Helper ─────────────────────────────────────────────────────────
def add_p(text, bold=False, size_emu=None, color=None, align=None,
          space_after=None, space_before=None, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size_emu:
        run.font.size = Emu(size_emu)
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Emu(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Emu(space_before)
    return p, run

def add_run_to(p, text, bold=False, size_emu=None, color=None, italic=False):
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size_emu:
        run.font.size = Emu(size_emu)
    if color:
        run.font.color.rgb = color
    return run

def shade_cells(row, color_hex):
    """Apply background shading to all cells in a row."""
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, bottom={"sz": 6, "color": "CCCCCC"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def make_borderless(table):
    """Remove all borders from table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_cell_text(cell, text, bold=False, size_emu=None, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if size_emu:
        run.font.size = Emu(size_emu)
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Emu(25400)
    p.paragraph_format.space_before = Emu(12700)

# ══════════════════════════════════════════════════════════════════
#  DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════

# ── Company header ─────────────────────────────────────────────────
add_p('Solum Invest GmbH', bold=True, size_emu=228600, color=C_TITLE, space_after=25400)
add_p('Augasse 4, 6111 Volders', size_emu=101600, color=C_ADDR, space_after=152400)

# ── Spacer ─────────────────────────────────────────────────────────
add_p('', space_after=76200)

# ── Header table: Recipient left, Invoice details right ────────────
htable = doc.add_table(rows=5, cols=2)
make_borderless(htable)
htable.columns[0].width = Cm(9)
htable.columns[1].width = Cm(7)

set_cell_text(htable.rows[0].cells[0], 'An:', size_emu=101600, color=C_DESC)
set_cell_text(htable.rows[0].cells[1], 'Rechnung Nr. 07/2026', bold=True, size_emu=133350, color=C_TITLE)
set_cell_text(htable.rows[1].cells[0], 'Time to Jump Dolomites', bold=True, size_emu=114300)
set_cell_text(htable.rows[1].cells[1], 'Datum: 28. April 2026', size_emu=101600, color=C_DESC)
set_cell_text(htable.rows[2].cells[0], 'David Prato', size_emu=114300, color=C_SECT)
set_cell_text(htable.rows[2].cells[1], 'Zahlungsstatus: Bezahlt', bold=True, size_emu=101600, color=RGBColor(0x2E, 0x7D, 0x32))
set_cell_text(htable.rows[3].cells[0], 'info@timetojumpdolomites.com', size_emu=101600, color=C_DESC)
set_cell_text(htable.rows[3].cells[1], '', size_emu=101600)
set_cell_text(htable.rows[4].cells[0], 'UID: ausstehend', size_emu=101600, color=C_DESC)
set_cell_text(htable.rows[4].cells[1], '', size_emu=101600)

# ── Title ──────────────────────────────────────────────────────────
add_p('RECHNUNG', bold=True, size_emu=254000, color=C_TITLE,
      align=WD_ALIGN_PARAGRAPH.LEFT, space_after=25400, space_before=177800)

add_p('Teilrechnung 1  ·  Projekt: Website Relaunch TimeToJump Dolomites',
      size_emu=101600, color=C_SUB, space_after=127000)

add_p('Referenz: Angebot Nr. 02/2026 vom 15. März 2026  ·  Gesamtauftragswert: EUR 900,–',
      size_emu=101600, color=C_SUB, space_after=127000)

# ── Intro text ─────────────────────────────────────────────────────
p, _ = add_p('', space_after=101600)
add_run_to(p, 'Sehr geehrter Herr Prato,', bold=True, size_emu=114300)

p, _ = add_p('', space_after=101600)
add_run_to(p, 'hiermit stellen wir Ihnen die erste Teilrechnung für die im Rahmen des Angebots Nr. 02/2026 erbrachten und abgenommenen Leistungen.', size_emu=114300, color=C_SECT)

# ── Section: Abgerechnete Leistungen ──────────────────────────────
add_p('Leistungsübersicht', bold=True, size_emu=152400, color=C_BLACK,
      space_before=127000, space_after=76200)

# Main table
table = doc.add_table(rows=1, cols=3)
make_borderless(table)
table.columns[0].width = Cm(1.2)
table.columns[1].width = Cm(11.5)
table.columns[2].width = Cm(3.3)

# Header row
set_cell_text(table.rows[0].cells[0], 'Pos.', bold=True, size_emu=101600, color=C_HEAD)
set_cell_text(table.rows[0].cells[1], 'Beschreibung', bold=True, size_emu=101600, color=C_HEAD)
set_cell_text(table.rows[0].cells[2], 'Status', bold=True, size_emu=101600, color=C_HEAD, align=WD_ALIGN_PARAGRAPH.RIGHT)
shade_cells(table.rows[0], 'F2F2F2')

items = [
    ('1', 'Vollständige Landingpage – 12 Sektionen, responsiv\n'
          'Hero, About, What to Expect, Pricing, Testimonials, FAQ,\n'
          'Anforderungen, Google Maps, Footer mit Social-Media-Links', '✓ Fertig'),
    ('2', 'Mehrsprachigkeit – 9 Sprachen mit LocalStorage\n'
          'EN, DE, IT, LAD, NL, FR, PL, ES, CS', '✓ Fertig'),
    ('3', 'Kontaktformular mit E-Mail-Versand\n'
          'FormSubmit.co, CAPTCHA + Honeypot, Fehlerbehandlung', '✓ Fertig'),
    ('4', 'WhatsApp / E-Mail / Telefon – schwebendes Chat-Widget\n'
          'Direktlinks, sichtbar auf allen Endgeräten', '✓ Fertig'),
    ('5', 'Sicherheit & SEO-Grundstruktur\n'
          'Security Headers (CSP), SSL/HTTPS, Meta-Tags, Open Graph,\n'
          'semantisches HTML5, Ladezeit < 1s', '✓ Fertig'),
    ('6', 'Datenschutzseite (privacy.html)', '✓ Fertig'),
    ('7', 'Feedback-Schleifen vor Go-Live (> 2 absolviert)', '✓ Fertig'),
]

for pos, desc, status in items:
    row = table.add_row()
    set_cell_text(row.cells[0], pos, bold=True, size_emu=114300, color=C_SECT)
    set_cell_text(row.cells[1], desc, size_emu=101600, color=C_SECT)
    set_cell_text(row.cells[2], status, size_emu=88900, color=RGBColor(0x2E, 0x7D, 0x32),
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

# ── Preistabelle ───────────────────────────────────────────────────
add_p('Preisübersicht', bold=True, size_emu=152400, color=C_BLACK,
      space_before=152400, space_after=76200)

ptable = doc.add_table(rows=4, cols=2)
make_borderless(ptable)
ptable.columns[0].width = Cm(12)
ptable.columns[1].width = Cm(4)

# Row 0: Teilbetrag
set_cell_text(ptable.rows[0].cells[0], 'Teilbetrag (Positionen 1–7)', size_emu=114300, color=C_SECT)
set_cell_text(ptable.rows[0].cells[1], 'EUR 600,–', bold=True, size_emu=114300, align=WD_ALIGN_PARAGRAPH.RIGHT)

# Row 1: USt
set_cell_text(ptable.rows[1].cells[0], 'Umsatzsteuer (Reverse Charge – Art. 196 MwSt-RL)', size_emu=101600, color=C_DESC, align=None)
set_cell_text(ptable.rows[1].cells[1], '0,00 EUR', size_emu=101600, color=C_DESC, align=WD_ALIGN_PARAGRAPH.RIGHT)

# Row 2: Separator visual (empty with bottom border)
set_cell_text(ptable.rows[2].cells[0], '', size_emu=50800)
set_cell_text(ptable.rows[2].cells[1], '', size_emu=50800)

# Row 3: Total
set_cell_text(ptable.rows[3].cells[0], 'Rechnungsbetrag', bold=True, size_emu=133350, color=C_TITLE)
set_cell_text(ptable.rows[3].cells[1], 'EUR 600,–', bold=True, size_emu=133350, color=C_TITLE, align=WD_ALIGN_PARAGRAPH.RIGHT)
shade_cells(ptable.rows[3], 'F2F2F2')

# ── Zahlungsstatus ─────────────────────────────────────────────────
add_p('Zahlungsstatus', bold=True, size_emu=127000, color=C_HEAD,
      space_before=203200, space_after=50800)

p, _ = add_p('', space_after=50800)
add_run_to(p, 'Der Rechnungsbetrag von EUR 600,– wurde bereits vollständig beglichen.',
           bold=True, size_emu=114300, color=RGBColor(0x2E, 0x7D, 0x32))

p, _ = add_p('', space_after=76200)
add_run_to(p, 'Verbleibender offener Betrag aus Angebot Nr. 02/2026:  ',
           size_emu=101600, color=C_DESC)
add_run_to(p, 'EUR 300,–', bold=True, size_emu=101600, color=C_SECT)
add_run_to(p, '  (Teilrechnung 2 – fällig nach Go-Live / DNS-Migration)',
           size_emu=101600, color=C_DESC)

# ── Steuerhinweis ──────────────────────────────────────────────────
add_p('Steuerhinweis', bold=True, size_emu=114300, color=C_HEAD,
      space_before=152400, space_after=50800)

add_p('Diese Leistung ist eine innergemeinschaftliche sonstige Leistung (B2B) gemäß '
      'Art. 196 MwSt-Richtlinie. Es wird keine österreichische Umsatzsteuer in Rechnung '
      'gestellt. Der Leistungsempfänger schuldet die Steuer im eigenen Land (Reverse Charge).',
      italic=True, size_emu=101600, color=C_DESC, space_after=127000)

# ── Signatur ───────────────────────────────────────────────────────
add_p('Mit freundlichen Grüßen,', size_emu=114300, color=C_SECT,
      space_before=152400, space_after=254000)

add_p('Axel Ludwig Mühlbacher', bold=True, size_emu=127000, color=C_TITLE, space_after=25400)
add_p('Geschäftsführer – Solum Invest GmbH', size_emu=101600, color=C_DESC, space_after=152400)

# ── Bankverbindung ─────────────────────────────────────────────────
btable = doc.add_table(rows=4, cols=2)
make_borderless(btable)
btable.columns[0].width = Cm(4)
btable.columns[1].width = Cm(12)

set_cell_text(btable.rows[0].cells[0], 'Bank:', size_emu=101600, color=C_DESC)
set_cell_text(btable.rows[0].cells[1], 'Raiffeisenbank Volders', size_emu=101600, color=C_SECT)
set_cell_text(btable.rows[1].cells[0], 'IBAN:', size_emu=101600, color=C_DESC)
set_cell_text(btable.rows[1].cells[1], 'AT08 3620 0000 0035 2971', size_emu=101600, color=C_SECT)
set_cell_text(btable.rows[2].cells[0], 'BIC:', size_emu=101600, color=C_DESC)
set_cell_text(btable.rows[2].cells[1], 'RZTIAT 22200', size_emu=101600, color=C_SECT)
set_cell_text(btable.rows[3].cells[0], 'Verwendung:', size_emu=101600, color=C_DESC)
set_cell_text(btable.rows[3].cells[1], 'Rechnung 07/2026 – TimeToJump Dolomites', size_emu=101600, color=C_SECT)

# ── Footer ─────────────────────────────────────────────────────────
add_p('Solum Invest GmbH  ·  Augasse 4, 6111 Volders  ·  ATU81665404  ·  FN 642566a  ·  office@soluminvest.at  ·  +43 676 4008771',
      size_emu=88900, color=C_ADDR, align=WD_ALIGN_PARAGRAPH.CENTER,
      space_before=254000, space_after=0)

# ══════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════
OUT_DIR = r"c:\AlpineDigital\timetojumpdolomites\Korrespondenz\Angebote und Rechnungen"
out_path = os.path.join(OUT_DIR, 'Rechnung_07_2026_Teilrechnung_TimeToJump_Dolomites.docx')
doc.save(out_path)
print(f'Saved: {out_path}')

# Also save a copy to the accounting folder
ACCT_DIR = r"G:\Andere Computer\Mein Laptop\Documents\01_Finanzen_und_Vermoegen\Buchhaltung\Buchhaltung Solum Invest GmbH\Buchhaltung GmbH 2026\04-April-2026"
acct_out = os.path.join(ACCT_DIR, 'Rechnungen Out - April 2026')
os.makedirs(acct_out, exist_ok=True)
acct_path = os.path.join(acct_out, 'Rechnung_07_2026_Teilrechnung_TimeToJump_Dolomites.docx')
doc.save(acct_path)
print(f'Saved: {acct_path}')
