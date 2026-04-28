"""
Rechnung 07/2026 – Final version with Solum Invest Corporate Design
(from Stockart Conversion Angebot: dark header, gold accents, Helvetica)
"""
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

# Corporate Design Colors
C_DARK    = RGBColor(0x23, 0x1F, 0x20)  # Dark anthracite - headers
C_GOLD    = RGBColor(0x8E, 0x79, 0x3E)  # Gold - accents, bullets, totals
C_GOLD_LT = RGBColor(0xAD, 0x97, 0x4F)  # Light gold - table headers
C_TEXT    = RGBColor(0x1A, 0x1A, 0x1A)  # Body text
C_GREY    = RGBColor(0x66, 0x66, 0x66)  # Secondary text
C_LIGHT   = RGBColor(0x99, 0x99, 0x99)  # Light text
C_GREEN   = RGBColor(0x2E, 0x7D, 0x32)  # Status green
C_AMBER   = RGBColor(0xC6, 0x8A, 0x00)  # Status amber
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

FILL_DARK   = '231F20'
FILL_ZEBRA  = 'F5F5F3'
FILL_HEADER = 'EBEBEB'

def add_p(text='', bold=False, sz=9, color=C_TEXT, align=None,
          sa=Pt(3), sb=Pt(0), italic=False, font='Helvetica'):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(sz); run.font.name = font
    if color: run.font.color.rgb = color
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = sa
    p.paragraph_format.space_before = sb
    return p

def add_run(p, text, bold=False, sz=9, color=C_TEXT, italic=False):
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(sz); run.font.name = 'Helvetica'
    if color: run.font.color.rgb = color
    return run

def shade(row, hex_color):
    for cell in row.cells:
        sh = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(sh)

def no_borders(table):
    tblPr = table._tbl.tblPr or parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        + ''.join(f'<w:{s} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                  for s in ['top','left','bottom','right','insideH','insideV'])
        + '</w:tblBorders>'))

def gold_bottom(row):
    """Add gold bottom border to row cells."""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="AE964F"/>'
            f'</w:tcBorders>'))

def sc(cell, text, bold=False, sz=9, color=C_TEXT, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(sz); run.font.name = 'Helvetica'
    if color: run.font.color.rgb = color
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

def scm(cell, lines, sz=8.5, color=C_TEXT, bold_first=True):
    cell.text = ''
    p = cell.paragraphs[0]
    for i, line in enumerate(lines):
        if i > 0: p.add_run('\n').font.size = Pt(sz)
        r = p.add_run(line)
        r.font.size = Pt(sz); r.font.name = 'Helvetica'; r.font.color.rgb = color
        if bold_first and i == 0: r.bold = True
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)

# ══════════════════════════════════════════════════════════════════
#  DARK HEADER BLOCK
# ══════════════════════════════════════════════════════════════════
ht = doc.add_table(rows=1, cols=2)
no_borders(ht)
shade(ht.rows[0], FILL_DARK)
ht.columns[0].width = Cm(9)
ht.columns[1].width = Cm(7)

# Left: Company
c0 = ht.rows[0].cells[0]
c0.text = ''
p = c0.paragraphs[0]
r = p.add_run('SOLUM INVEST GMBH')
r.bold = True; r.font.size = Pt(16); r.font.name = 'Helvetica'; r.font.color.rgb = C_WHITE
p.paragraph_format.space_after = Pt(2); p.paragraph_format.space_before = Pt(8)
p2 = c0.add_paragraph()
r2 = p2.add_run('Augasse 4, 6111 Volders  \u00b7  +43 676 4008771')
r2.font.size = Pt(8); r2.font.name = 'Helvetica'; r2.font.color.rgb = RGBColor(0xBB,0xBB,0xBB)
p2.paragraph_format.space_after = Pt(1)
p3 = c0.add_paragraph()
r3 = p3.add_run('office@soluminvest.at  \u00b7  ATU81665404  \u00b7  FN 642566a')
r3.font.size = Pt(8); r3.font.name = 'Helvetica'; r3.font.color.rgb = RGBColor(0xBB,0xBB,0xBB)
p3.paragraph_format.space_after = Pt(8)

# Right: Invoice info
c1 = ht.rows[0].cells[1]
c1.text = ''
p = c1.paragraphs[0]
r = p.add_run('RECHNUNG')
r.bold = True; r.font.size = Pt(18); r.font.name = 'Helvetica'; r.font.color.rgb = C_GOLD
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_after = Pt(4); p.paragraph_format.space_before = Pt(8)
p2 = c1.add_paragraph()
r2 = p2.add_run('Nr. 07/2026')
r2.bold = True; r2.font.size = Pt(11); r2.font.name = 'Helvetica'; r2.font.color.rgb = C_WHITE
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p2.paragraph_format.space_after = Pt(1)
p3 = c1.add_paragraph()
r3 = p3.add_run('Datum: 28. April 2026')
r3.font.size = Pt(8); r3.font.name = 'Helvetica'; r3.font.color.rgb = RGBColor(0xBB,0xBB,0xBB)
p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p3.paragraph_format.space_after = Pt(8)

# ── Gold separator line ────────────────────────────────────────────
sep = doc.add_table(rows=1, cols=1)
no_borders(sep)
gold_bottom(sep.rows[0])
sep.rows[0].cells[0].text = ''
sep.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)
sep.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(0)

# ── Recipient + Subtitle ──────────────────────────────────────────
add_p('', sa=Pt(6))

rt = doc.add_table(rows=4, cols=2)
no_borders(rt)
rt.columns[0].width = Cm(8.5)
rt.columns[1].width = Cm(7.5)

sc(rt.rows[0].cells[0], 'An:', sz=8, color=C_GREY)
sc(rt.rows[0].cells[1], 'Teilrechnung 1', bold=True, sz=9, color=C_GOLD, align=WD_ALIGN_PARAGRAPH.RIGHT)
sc(rt.rows[1].cells[0], 'Time to Jump Dolomites', bold=True, sz=10, color=C_DARK)
sc(rt.rows[1].cells[1], 'Projekt: Website Relaunch TimeToJump Dolomites', sz=8, color=C_GREY, align=WD_ALIGN_PARAGRAPH.RIGHT)
sc(rt.rows[2].cells[0], 'David Prato  \u00b7  info@timetojumpdolomites.com', sz=8.5, color=C_GREY)
sc(rt.rows[2].cells[1], 'Ref.: Angebot Nr. 02/2026 vom 15. M\u00e4rz 2026', sz=8, color=C_GREY, align=WD_ALIGN_PARAGRAPH.RIGHT)
sc(rt.rows[3].cells[0], 'UID: IT03113270213', sz=8.5, color=C_GREY)
sc(rt.rows[3].cells[1], 'Gesamtauftragswert: EUR 2.000,\u2013', bold=True, sz=8.5, color=C_DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)

# ── Intro ──────────────────────────────────────────────────────────
add_p('', sa=Pt(4))
p = add_p('', sa=Pt(6))
add_run(p, 'Sehr geehrter Herr Prato, ', bold=True, sz=9)
add_run(p, 'hiermit stellen wir Ihnen die erste Teilrechnung f\u00fcr die im Rahmen des Angebots Nr. 02/2026 sowie der telefonischen Vereinbarung vom 27.04.2026 erbrachten und beauftragten Leistungen.', sz=9, color=C_TEXT)

# ══════════════════════════════════════════════════════════════════
#  SECTION A: Angebot 02 Leistungen
# ══════════════════════════════════════════════════════════════════
add_p('1. Angebot Nr. 02/2026 \u2013 Website Relaunch', bold=True, sz=13, color=C_DARK, sb=Pt(10), sa=Pt(6))

t1 = doc.add_table(rows=1, cols=3)
no_borders(t1)
t1.columns[0].width = Cm(0.9)
t1.columns[1].width = Cm(12)
t1.columns[2].width = Cm(3.1)

# Gold header row
sc(t1.rows[0].cells[0], 'Pos.', bold=True, sz=8.5, color=C_GOLD_LT)
sc(t1.rows[0].cells[1], 'Beschreibung', bold=True, sz=8.5, color=C_GOLD_LT)
sc(t1.rows[0].cells[2], 'Status', bold=True, sz=8.5, color=C_GOLD_LT, align=WD_ALIGN_PARAGRAPH.RIGHT)
gold_bottom(t1.rows[0])

items_a = [
    ('1', ['Vollst\u00e4ndige Landingpage \u2013 12 Sektionen, responsiv',
           'Hero, About, What to Expect, Pricing, Testimonials, FAQ, Anforderungen, Google Maps, Footer']),
    ('2', ['Mehrsprachigkeit \u2013 9 Sprachen mit LocalStorage',
           'EN, DE, IT, LAD, NL, FR, PL, ES, CS']),
    ('3', ['Kontaktformular mit E-Mail-Versand',
           'FormSubmit.co, CAPTCHA + Honeypot, Fehlerbehandlung']),
    ('4', ['WhatsApp / E-Mail / Telefon \u2013 schwebendes Chat-Widget']),
    ('5', ['Sicherheit & SEO-Grundstruktur',
           'Security Headers, SSL/HTTPS, Meta-Tags, Open Graph, semantisches HTML5']),
    ('6', ['Datenschutzseite (privacy.html)']),
    ('7', ['Feedback-Schleifen vor Go-Live (> 2 absolviert)']),
]

for i, (pos, lines) in enumerate(items_a):
    row = t1.add_row()
    sc(row.cells[0], pos, bold=True, sz=9, color=C_GOLD)
    scm(row.cells[1], lines, sz=8.5, color=C_TEXT)
    sc(row.cells[2], 'Fertig', sz=8.5, color=C_GREEN, align=WD_ALIGN_PARAGRAPH.RIGHT)
    if i % 2 == 1: shade(row, FILL_ZEBRA)

# ══════════════════════════════════════════════════════════════════
#  SECTION B: Telefonische Absprache
# ══════════════════════════════════════════════════════════════════
add_p('2. Telefonische Absprache vom 27.04.2026 \u2013 Erweiterungen', bold=True, sz=13, color=C_DARK, sb=Pt(10), sa=Pt(6))

t2 = doc.add_table(rows=1, cols=3)
no_borders(t2)
t2.columns[0].width = Cm(0.9)
t2.columns[1].width = Cm(12)
t2.columns[2].width = Cm(3.1)

sc(t2.rows[0].cells[0], 'Pos.', bold=True, sz=8.5, color=C_GOLD_LT)
sc(t2.rows[0].cells[1], 'Beschreibung', bold=True, sz=8.5, color=C_GOLD_LT)
sc(t2.rows[0].cells[2], 'Status', bold=True, sz=8.5, color=C_GOLD_LT, align=WD_ALIGN_PARAGRAPH.RIGHT)
gold_bottom(t2.rows[0])

items_b = [
    ('8',  ['Supabase-Backend mit PostgreSQL-Datenbank',
            'Event- & Buchungstabellen, Row Level Security, Edge Functions']),
    ('9',  ['Booking Engine & dynamische Flight Cards',
            'Live-Verf\u00fcgbarkeit, sequenzielle Load-Logik, Nickname-Pills']),
    ('10', ['Dynamisches Gruppenpricing',
            'Preis pro Person sinkt automatisch mit jedem weiteren Springer']),
    ('11', ['Buchungsformular (Booking Modal)',
            'Alle Pflichtfelder, Glassmorphism-Design, vollst\u00e4ndig in 9 Sprachen']),
    ('12', ['Stripe-Integration (Kreditkarten-Hold)',
            'SetupIntent, 0\u20ac Card Hold, Strafeinzug bei Stornierung']),
    ('13', ['AGB & Rechtstexte',
            'Stornobedingungen (7-Tage-Regel), Briefing-Strafe, Zero Tolerance']),
    ('14', ['Admin-Dashboard (admin.html)',
            'Event-Verwaltung & Buchungs\u00fcbersicht']),
    ('15', ['Website-Optimierungen & Content-Updates',
            'Pricing, FAQ, Maps, Navigation, Elikos-Bereinigung']),
    ('16', ['Erfolgsseite & WhatsApp-Gruppeneinladung']),
    ('17', ['DNS-Migration & Wix-Abl\u00f6se',
            'Domain-Umzug auf Cloudflare Pages, SSL, Weiterleitungen']),
    ('18', ['Automatische Reminder-E-Mails & Gruppenbuchung',
            '7/5/3 Tage vor Abflug, Hauptperson bucht f\u00fcr Gruppe']),
]

for i, (pos, lines) in enumerate(items_b):
    row = t2.add_row()
    sc(row.cells[0], pos, bold=True, sz=9, color=C_GOLD)
    scm(row.cells[1], lines, sz=8.5, color=C_TEXT)
    sc(row.cells[2], 'In Bearbeitung', sz=8.5, color=C_AMBER, align=WD_ALIGN_PARAGRAPH.RIGHT)
    if i % 2 == 1: shade(row, FILL_ZEBRA)

# ══════════════════════════════════════════════════════════════════
#  PREISTABELLE
# ══════════════════════════════════════════════════════════════════
add_p('3. Preis\u00fcbersicht', bold=True, sz=13, color=C_DARK, sb=Pt(12), sa=Pt(6))

pt = doc.add_table(rows=5, cols=2)
no_borders(pt)
pt.columns[0].width = Cm(12)
pt.columns[1].width = Cm(4)

sc(pt.rows[0].cells[0], 'Pos.', bold=True, sz=8.5, color=C_GOLD_LT)
sc(pt.rows[0].cells[1], 'Betrag netto', bold=True, sz=8.5, color=C_GOLD_LT, align=WD_ALIGN_PARAGRAPH.RIGHT)
gold_bottom(pt.rows[0])

sc(pt.rows[1].cells[0], 'Teilbetrag Positionen 1\u20137 (Angebot Nr. 02/2026)', sz=9, color=C_TEXT)
sc(pt.rows[1].cells[1], 'EUR 600,\u2013', bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

sc(pt.rows[2].cells[0], 'Umsatzsteuer (Reverse Charge \u2013 Art. 196 MwSt-RL)', sz=8, color=C_GREY)
sc(pt.rows[2].cells[1], '0,00 EUR', sz=8, color=C_GREY, align=WD_ALIGN_PARAGRAPH.RIGHT)

sc(pt.rows[3].cells[0], '', sz=4)
sc(pt.rows[3].cells[1], '', sz=4)

# Total row with gold
sc(pt.rows[4].cells[0], 'Rechnungsbetrag netto', bold=True, sz=9.5, color=C_GOLD)
sc(pt.rows[4].cells[1], 'EUR 600,\u2013', bold=True, sz=9.5, color=C_GOLD, align=WD_ALIGN_PARAGRAPH.RIGHT)
gold_bottom(pt.rows[4])

# ── Tax note ───────────────────────────────────────────────────────
add_p('Steuerhinweis: Innergemeinschaftliche sonstige Leistung (B2B) gem. Art. 196 MwSt-Richtlinie. '
      'Es wird keine \u00f6sterreichische USt in Rechnung gestellt. Der Leistungsempf\u00e4nger schuldet die '
      'Steuer im eigenen Land (Reverse Charge). Bitte Ihre UID-Nr. im Aufnahmefeld eintragen.',
      italic=True, sz=8, color=C_GREY, sa=Pt(10))

# ══════════════════════════════════════════════════════════════════
#  ZAHLUNGSSTATUS & OFFENE BETRAEGE
# ══════════════════════════════════════════════════════════════════
add_p('4. Zahlungsstatus & offene Betr\u00e4ge', bold=True, sz=13, color=C_DARK, sb=Pt(8), sa=Pt(6))

p = add_p('', sa=Pt(8))
add_run(p, '\u25cf  ', bold=True, sz=9, color=C_GREEN)
add_run(p, 'Der Rechnungsbetrag von EUR 600,\u2013 wurde bereits vollst\u00e4ndig beglichen.', bold=True, sz=9, color=C_GREEN)

ot = doc.add_table(rows=4, cols=3)
no_borders(ot)
ot.columns[0].width = Cm(8.5)
ot.columns[1].width = Cm(3.5)
ot.columns[2].width = Cm(4)

sc(ot.rows[0].cells[0], '', bold=True, sz=8.5, color=C_GOLD_LT)
sc(ot.rows[0].cells[1], 'Betrag', bold=True, sz=8.5, color=C_GOLD_LT, align=WD_ALIGN_PARAGRAPH.RIGHT)
sc(ot.rows[0].cells[2], 'F\u00e4lligkeit', bold=True, sz=8.5, color=C_GOLD_LT, align=WD_ALIGN_PARAGRAPH.RIGHT)
gold_bottom(ot.rows[0])

sc(ot.rows[1].cells[0], 'Verbleibender offener Betrag Angebot Nr. 02/2026', sz=9, color=C_TEXT)
sc(ot.rows[1].cells[1], 'EUR 300,\u2013', bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
sc(ot.rows[1].cells[2], 'F\u00e4llig Ende Mai', sz=8.5, color=C_AMBER, align=WD_ALIGN_PARAGRAPH.RIGHT)

sc(ot.rows[2].cells[0], 'Verbleibender offener Betrag aus\ntelefonischer Absprache 27.04.2026', sz=9, color=C_TEXT)
sc(ot.rows[2].cells[1], 'EUR 1.100,\u2013', bold=True, sz=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
sc(ot.rows[2].cells[2], 'F\u00e4llig nach Go-Live\nmit DNS-Umstellung', sz=8.5, color=C_AMBER, align=WD_ALIGN_PARAGRAPH.RIGHT)
shade(ot.rows[2], FILL_ZEBRA)

sc(ot.rows[3].cells[0], 'Gesamt offen', bold=True, sz=9.5, color=C_GOLD)
sc(ot.rows[3].cells[1], 'EUR 1.400,\u2013', bold=True, sz=9.5, color=C_GOLD, align=WD_ALIGN_PARAGRAPH.RIGHT)
sc(ot.rows[3].cells[2], '', sz=8.5)
gold_bottom(ot.rows[3])

# ══════════════════════════════════════════════════════════════════
#  SIGNATUR
# ══════════════════════════════════════════════════════════════════
add_p('', sa=Pt(4))
p = add_p('', sa=Pt(6))
add_run(p, 'Zahlungsbedingungen: ', bold=True, sz=9)
add_run(p, 'Zahlungsziel 14 Tage netto. Die offenen Betr\u00e4ge werden gem\u00e4\u00df den oben genannten F\u00e4lligkeiten in Rechnung gestellt.', sz=9)

add_p('Mit freundlichen Gr\u00fc\u00dfen,', sz=9, sa=Pt(20), sb=Pt(10))
add_p('Axel Ludwig M\u00fchlbacher', bold=True, sz=9, color=C_TEXT, sa=Pt(1))
add_p('Gesch\u00e4ftsf\u00fchrer \u2013 Solum Invest GmbH', sz=9, color=C_TEXT, sa=Pt(8))

# ══════════════════════════════════════════════════════════════════
#  FOOTER: Bank + Gold line + Company
# ══════════════════════════════════════════════════════════════════
# Gold separator
sep2 = doc.add_table(rows=1, cols=1)
no_borders(sep2)
gold_bottom(sep2.rows[0])
sep2.rows[0].cells[0].text = ''
sep2.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)

add_p('', sa=Pt(2))

# Bank
p = add_p('', sa=Pt(2))
add_run(p, 'Bankverbindung', bold=True, sz=8.5, color=C_GOLD)

p = add_p('', sa=Pt(1))
add_run(p, 'Solum Invest GmbH  \u00b7  Raiffeisenbank Volders', sz=8, color=C_TEXT)

p = add_p('', sa=Pt(1))
add_run(p, 'IBAN: AT08 3620 0000 0035 2971  \u00b7  BIC: RZTIAT 22200', sz=8, color=C_TEXT)

p = add_p('', sa=Pt(6))
add_run(p, 'Verwendungszweck: Rechnung 07/2026 \u2013 TimeToJump Dolomites', sz=8, color=C_TEXT)

# Footer line
add_p('Solum Invest GmbH  \u00b7  Augasse 4  \u00b7  6111 Volders  \u00b7  ATU81665404  \u00b7  FN 642566a  \u00b7  office@soluminvest.at  \u00b7  +43 676 4008771',
      sz=7, color=C_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, sb=Pt(8))

# ══════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════
FN = 'Rechnung_07_2026_Teilrechnung_TimeToJump_Dolomites.docx'

out1 = os.path.join(r'c:\AlpineDigital\timetojumpdolomites\Korrespondenz\Angebote und Rechnungen', FN)
doc.save(out1)
print(f'Saved: {out1}')

acct = r'G:\Andere Computer\Mein Laptop\Documents\01_Finanzen_und_Vermoegen\Buchhaltung\Buchhaltung Solum Invest GmbH\Buchhaltung GmbH 2026\04-April-2026\Rechnungen Out - April 2026'
os.makedirs(acct, exist_ok=True)
doc.save(os.path.join(acct, FN))
print(f'Saved: {os.path.join(acct, FN)}')
