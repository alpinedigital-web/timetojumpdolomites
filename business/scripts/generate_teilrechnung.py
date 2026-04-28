from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

def add_styled_paragraph(text, bold=False, size=10, color=None, alignment=None, space_after=Pt(4), space_before=Pt(0)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    return p

def set_cell_text(cell, text, bold=False, size=10, alignment=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if alignment:
        p.alignment = alignment

# --- HEADER ---
add_styled_paragraph('Solum Invest GmbH', bold=True, size=14, space_after=Pt(2))
add_styled_paragraph('Augasse 4 · 6111 Volders · Österreich', size=9, space_after=Pt(1))
add_styled_paragraph('Tel.: +43 676 4008771  ·  office@soluminvest.at', size=9, space_after=Pt(1))
add_styled_paragraph('UID: ATU81665404  ·  FN 642566a (LG Innsbruck)', size=9, space_after=Pt(12))

# --- INVOICE TITLE ---
add_styled_paragraph('TEILRECHNUNG Nr. 01/2026', bold=True, size=16, space_after=Pt(4))
add_styled_paragraph('Datum: 28. April 2026', size=10, space_after=Pt(2))
add_styled_paragraph('Referenz: Angebot Nr. 02/2026 vom 15. März 2026', size=10, space_after=Pt(12))

# --- RECIPIENT ---
add_styled_paragraph('An:', bold=True, size=10, space_after=Pt(2))
add_styled_paragraph('Time to Jump Dolomites', bold=True, size=11, space_after=Pt(1))
add_styled_paragraph('David Prato', size=10, space_after=Pt(1))
add_styled_paragraph('www.timetojumpdolomites.com', size=10, space_after=Pt(1))
add_styled_paragraph('info@timetojumpdolomites.com', size=10, space_after=Pt(12))

# --- INTRO ---
add_styled_paragraph('Sehr geehrter Herr Prato,', bold=True, size=10, space_after=Pt(8))

p = doc.add_paragraph()
run = p.add_run('hiermit stellen wir Ihnen die erste Teilrechnung für die im Rahmen des Angebots Nr. 02/2026 erbrachten Leistungen. Diese Teilrechnung deckt die bereits fertiggestellten und abgenommenen Leistungen gemäß Leistungsumfang §1 des Angebots ab.')
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(12)

# --- LEISTUNGEN TABLE ---
add_styled_paragraph('Abgerechnete Leistungen', bold=True, size=12, space_after=Pt(6))

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
headers = ['Pos.', 'Beschreibung', 'Betrag netto']
for i, h in enumerate(headers):
    set_cell_text(table.rows[0].cells[i], h, bold=True, size=10)

# Set column widths
table.columns[0].width = Cm(1.5)
table.columns[1].width = Cm(11)
table.columns[2].width = Cm(3.5)

rows_data = [
    ('1.', 'Vollständige Landingpage (12 Sektionen, responsiv – Desktop, Tablet, Smartphone)', ''),
    ('', '  – Hero-Bereich mit Hintergrundbild und Call-to-Action', ''),
    ('', '  – About-Sektion mit Statistiken', ''),
    ('', '  – What to Expect (4-Schritte-Ablauf)', ''),
    ('', '  – Pricing (3 Preiskarten mit CTAs)', ''),
    ('', '  – Testimonials, FAQ, Anforderungen', ''),
    ('', '  – Google Maps, Footer mit Social-Media-Links', ''),
    ('2.', 'Mehrsprachigkeit – 9 Sprachen (EN, DE, IT, LAD, NL, FR, PL, ES, CS) mit LocalStorage', ''),
    ('3.', 'Kontaktformular mit E-Mail-Versand (FormSubmit.co, CAPTCHA, Honeypot)', ''),
    ('4.', 'WhatsApp / E-Mail / Telefon – schwebendes Chat-Widget', ''),
    ('5.', 'Sicherheit & SEO-Grundstruktur (Security Headers, Meta-Tags, SSL, semantisches HTML)', ''),
    ('6.', 'Datenschutzseite (privacy.html)', ''),
    ('7.', '2 Feedback-Schleifen vor Go-Live (mehr als 2 absolviert)', ''),
    ('', '', ''),
    ('', 'Teilbetrag für obige Leistungen (Teilrechnung 1 von Gesamtbetrag EUR 900,–)', 'EUR 600,–'),
]

for pos, desc, amount in rows_data:
    row = table.add_row()
    set_cell_text(row.cells[0], pos, bold=bool(pos and pos[0].isdigit()), size=10)
    set_cell_text(row.cells[1], desc, bold=('Teilbetrag' in desc), size=10)
    set_cell_text(row.cells[2], amount, bold=bool(amount), size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

# Summary rows
for label, amount, bold in [
    ('Umsatzsteuer (Reverse Charge – Art. 196 MwSt-RL)', '0,00 EUR', False),
    ('Rechnungsbetrag netto', 'EUR 600,–', True),
]:
    row = table.add_row()
    set_cell_text(row.cells[0], '', size=10)
    set_cell_text(row.cells[1], label, bold=bold, size=10)
    set_cell_text(row.cells[2], amount, bold=bold, size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_paragraph()

# --- PAYMENT NOTE ---
add_styled_paragraph('Zahlungsstatus', bold=True, size=12, space_after=Pt(6))

p = doc.add_paragraph()
run = p.add_run('Der Rechnungsbetrag von EUR 600,– wurde bereits vollständig beglichen. Wir bedanken uns für die pünktliche Zahlung.')
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run('Verbleibender offener Betrag aus Angebot Nr. 02/2026: EUR 300,– (Teilrechnung 2 nach Go-Live / DNS-Migration).')
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(12)

# --- TAX NOTE ---
add_styled_paragraph('Steuerhinweis', bold=True, size=11, space_after=Pt(4))
p = doc.add_paragraph()
run = p.add_run('Diese Leistung ist eine innergemeinschaftliche sonstige Leistung (B2B) gemäß Art. 196 MwSt-Richtlinie. Es wird keine österreichische Umsatzsteuer in Rechnung gestellt. Der Leistungsempfänger schuldet die Steuer im eigenen Land (Reverse Charge).')
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.italic = True
p.paragraph_format.space_after = Pt(12)

# --- SIGNATURE ---
add_styled_paragraph('Mit freundlichen Grüßen,', size=10, space_after=Pt(20))
add_styled_paragraph('Axel Ludwig Mühlbacher', bold=True, size=11, space_after=Pt(2))
add_styled_paragraph('Geschäftsführer – Solum Invest GmbH', size=10, space_after=Pt(12))

# --- BANK DETAILS ---
add_styled_paragraph('Bankverbindung', bold=True, size=10, space_after=Pt(4))
add_styled_paragraph('Solum Invest GmbH', size=9, space_after=Pt(1))
add_styled_paragraph('Raiffeisenbank Volders', size=9, space_after=Pt(1))
add_styled_paragraph('IBAN: AT08 3620 0000 0035 2971', size=9, space_after=Pt(1))
add_styled_paragraph('BIC: RZTIAT 22200', size=9, space_after=Pt(1))
add_styled_paragraph('Verwendungszweck: Teilrechnung 01/2026 – TimeToJump Dolomites', size=9, space_after=Pt(8))

# --- FOOTER ---
add_styled_paragraph('Solum Invest GmbH · Augasse 4 · 6111 Volders · ATU81665404 · FN 642566a · office@soluminvest.at · +43 676 4008771', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=(128,128,128))

output_path = os.path.join(os.path.dirname(__file__), 'Teilrechnung_01_2026_TimeToJump_Dolomites.docx')
doc.save(output_path)
print(f'Teilrechnung saved to: {output_path}')
