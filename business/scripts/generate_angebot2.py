from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

def add_p(text, bold=False, size=10, color=None, align=None, after=Pt(4), before=Pt(0), italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color: run.font.color.rgb = RGBColor(*color)
    if align: p.alignment = align
    p.paragraph_format.space_after = after
    p.paragraph_format.space_before = before
    return p

def set_cell(cell, text, bold=False, size=10, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if align: p.alignment = align

# HEADER
add_p('Solum Invest GmbH', bold=True, size=14, after=Pt(2))
add_p('Augasse 4 · 6111 Volders · Österreich', size=9, after=Pt(1))
add_p('Tel.: +43 676 4008771  ·  office@soluminvest.at', size=9, after=Pt(1))
add_p('UID: ATU81665404  ·  FN 642566a (LG Innsbruck)', size=9, after=Pt(12))

# TITLE
add_p('ANGEBOT Nr. 03/2026', bold=True, size=16, after=Pt(4))
add_p('Datum: 28. April 2026', size=10, after=Pt(2))
add_p('Gültig bis: 28. Mai 2026', size=10, after=Pt(12))

# RECIPIENT
add_p('An:', bold=True, after=Pt(2))
add_p('Time to Jump Dolomites', bold=True, size=11, after=Pt(1))
add_p('David Prato', after=Pt(1))
add_p('www.timetojumpdolomites.com', after=Pt(1))
add_p('info@timetojumpdolomites.com', after=Pt(8))

add_p('Betreff: Erweiterung Buchungsplattform – Backend, Booking Engine & Optimierungen', bold=True, size=11, after=Pt(4))
add_p('Ansprechpartner: Axel L. Mühlbacher, GF – Solum Invest GmbH', after=Pt(12))

# INTRO
add_p('Sehr geehrter Herr Prato,', bold=True, after=Pt(8))

p = doc.add_paragraph()
run = p.add_run('basierend auf unserem Vor-Ort-Meeting am 7. April 2026 sowie dem Telefonat am 27. April 2026 erhalten Sie hiermit unser Angebot für die Erweiterung Ihrer Website um eine vollständige Buchungsplattform inkl. Backend-Infrastruktur, Zahlungsintegration und weiterer besprochener Funktionen.')
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
run = p.add_run('Hinweis: Die nachfolgend aufgeführten Leistungen gehen über den Umfang des ursprünglichen Angebots Nr. 02/2026 (reine Website-Entwicklung, EUR 900,–) hinaus. Ein wesentlicher Teil dieser Leistungen wurde bereits als Goodwill vorab umgesetzt.')
run.font.size = Pt(10)
run.font.name = 'Calibri'
run.italic = True
p.paragraph_format.space_after = Pt(12)

# --- 1. LEISTUNGSUMFANG ---
add_p('1. Leistungsumfang', bold=True, size=13, after=Pt(8))

add_p('1.1  Backend & Datenbank (Supabase)', bold=True, size=11, after=Pt(4))
items_1_1 = [
    'Einrichtung Supabase-Backend mit PostgreSQL-Datenbank',
    'Datenstruktur für Events und Bookings inkl. Row Level Security (RLS)',
    'Admin-Dashboard (admin.html) zur Verwaltung von Events und Buchungen',
    'Edge Functions für serverseitige Geschäftslogik',
]
for item in items_1_1:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

add_p('1.2  Booking Engine & Flight Cards', bold=True, size=11, after=Pt(4), before=Pt(8))
items_1_2 = [
    'Dynamische Flight Cards mit Live-Verfügbarkeit aus der Datenbank',
    'Sequenzielle Load-Logik: Load 2 erst sichtbar/buchbar, wenn Load 1 voll ist',
    'Nickname-Pills: Öffentliche Passagierliste als Social Proof',
    'Dynamisches Gruppenpricing: Preis pro Person sinkt mit jedem weiteren Springer',
    'Beide Sprunggebiete (Mont de Côi & Saslong) gleichzeitig buchbar',
    'Kapazitätsmanagement: bis zu 10+10 Sprünge pro Tag',
]
for item in items_1_2:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

add_p('1.3  Buchungsformular (Booking Modal)', bold=True, size=11, after=Pt(4), before=Pt(8))
items_1_3 = [
    'Glassmorphism-Design, vollständig responsiv',
    'Alle Pflichtfelder: Vorname, Nachname, E-Mail, Telefon, vollständige Adresse, Geburtsdatum & -ort',
    'Fallschirmspringerlizenz (Land, Nummer, min. 100 Sprünge)',
    'Versicherungsnachweis (mit Verweis auf pianetavolo.it)',
    'Notfallkontakt (Name + Telefonnummer, getrennte Felder)',
    'Nickname für öffentliche Passagierliste',
    'AGB-Bestätigung mit Verlinkung',
    'Vollständige Lokalisierung in 9 Sprachen',
]
for item in items_1_3:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

add_p('1.4  Stripe-Integration (Kreditkarten-Hold)', bold=True, size=11, after=Pt(4), before=Pt(8))
items_1_4 = [
    'Stripe SetupIntent-Integration für 0€ Card Hold bei Reservierung',
    'Edge Function für sichere serverseitige Abwicklung',
    'Strafeinzug bei Stornierung oder verpasstem Briefing möglich',
    'Freigabe der Karte nach erfolgreichem Sprung',
    'Warnhinweis bei kurzfristiger Buchung (< 7 Tage): keine Erstattung',
]
for item in items_1_4:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

add_p('1.5  AGB & Rechtstexte', bold=True, size=11, after=Pt(4), before=Pt(8))
items_1_5 = [
    'AGB-Seite (terms.html) mit Stornobedingungen (7-Tage-Regel, 30% / 100%)',
    'Briefing-Strafe (50 €) bei Nichterscheinen am Freitagabend',
    'Stornierung durch Veranstalter bis 5 Tage vorher (volle Erstattung)',
    'Alkohol- & Drogenverbot klar kommuniziert (Zero Tolerance)',
    'Reservierung ausschließlich über Buchungsplattform (kein WhatsApp/E-Mail/Telefon)',
    'Headings in 9 Sprachen lokalisiert',
]
for item in items_1_5:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

add_p('1.6  Website-Optimierungen & Content-Updates', bold=True, size=11, after=Pt(4), before=Pt(8))
items_1_6 = [
    'Pricing-Bereinigung: 10er Tickets entfernt, Equipment Rental vereinfacht (100€ + 80€/Tag)',
    'Konkrete Gruppenpreis-Beträge entfernt, flexible Formulierung',
    'FAQ-Korrektur: Elikos-Buchungshinweis entfernt/korrigiert',
    'Standorte Mont de Côi + Saslong korrekt in Google Maps',
    'Partner-Referenzen (Elikos) bereinigt – Buchung nur über Time to Jump Dolomites',
    'Navigation bereinigt: Shop, Blog, leere Sektionen entfernt',
    'Erfolgsseite (success.html) mit WhatsApp-Gruppeneinladung',
    'Kontextuelle Tooltip-Hints bei kritischen Formularfeldern',
    'Flight Card Visual Fixes und Kontrastoptimierung',
]
for item in items_1_6:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

add_p('1.7  Noch umzusetzende Leistungen (in diesem Angebot enthalten)', bold=True, size=11, after=Pt(4), before=Pt(8))
items_1_7 = [
    'Anzahlung dynamisch = tatsächlicher Helikopterpreis (nicht pauschal)',
    'Gruppenbuchung: Hauptperson bucht, Formular per E-Mail an Mitspringer',
    'Automatische Reminder-E-Mails (7 / 5 / 3 Tage vor Abflug)',
    'Stripe Live-Checkout (Karte tatsächlich blocken im Produktivmodus)',
    'DNS-Migration & Wix-Ablöse (Domain-Umzug auf Cloudflare Pages)',
    'Preisanpassungen nach Bestätigung der Treibstoffkosten durch David',
]
for item in items_1_7:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

# --- 2. MITWIRKUNGSPFLICHTEN ---
add_p('2. Mitwirkungspflichten des Kunden', bold=True, size=13, after=Pt(6), before=Pt(12))
items_2 = [
    'Aruba DNS-Zugangsdaten (Login oder API-Token) für Domain-Umzug',
    'Stripe Business Account einrichten und API-Keys bereitstellen',
    'FormSubmit.co Bestätigungs-E-Mail einmalig bestätigen',
    'Bestätigung der Helikopterpreise 2026 (nach Treibstoffkostenprüfung)',
    'GMB-Validierung: Büro-Foto für Google My Business',
    'Eigene Fotos (2 Stück) per WhatsApp bereitstellen',
    'Finale Prüfung der italienischen und ladinischen Übersetzung',
]
for item in items_2:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)

# --- 3. PREISÜBERSICHT ---
add_p('3. Preisübersicht', bold=True, size=13, after=Pt(6), before=Pt(12))

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.columns[0].width = Cm(1.5)
table.columns[1].width = Cm(11)
table.columns[2].width = Cm(3.5)

for i, h in enumerate(['Pos.', 'Beschreibung', 'Betrag netto']):
    set_cell(table.rows[0].cells[i], h, bold=True)

price_rows = [
    ('1.', 'Backend & Datenbank (Supabase, RLS, Admin-Dashboard, Edge Functions)', ''),
    ('2.', 'Booking Engine (Flight Cards, Load-Logik, Gruppenpricing, Nickname-Pills)', ''),
    ('3.', 'Buchungsformular (Booking Modal, 9 Sprachen, alle Pflichtfelder)', ''),
    ('4.', 'Stripe-Integration (SetupIntent, Card Hold, Strafeinzug)', ''),
    ('5.', 'AGB & Rechtstexte (Storno, Briefing-Strafe, Zero Tolerance)', ''),
    ('6.', 'Website-Optimierungen & Content-Updates (Pricing, FAQ, Maps, Navigation)', ''),
    ('7.', 'Automatische Reminder-E-Mails, Gruppenbuchung, DNS-Migration', ''),
    ('', '', ''),
    ('', 'Gesamtbetrag netto', 'EUR 2.000,–'),
]
for pos, desc, amt in price_rows:
    row = table.add_row()
    set_cell(row.cells[0], pos, bold=bool(pos), size=10)
    set_cell(row.cells[1], desc, bold=('Gesamt' in desc), size=10)
    set_cell(row.cells[2], amt, bold=bool(amt), size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

for label, amount, b in [
    ('Umsatzsteuer (Reverse Charge – Art. 196 MwSt-RL)', '0,00 EUR', False),
    ('Rechnungsbetrag netto (= Gesamtbetrag)', 'EUR 2.000,–', True),
]:
    row = table.add_row()
    set_cell(row.cells[0], '')
    set_cell(row.cells[1], label, bold=b)
    set_cell(row.cells[2], amount, bold=b, align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_paragraph()

# Tax note
add_p('Steuerhinweis: Diese Leistung ist eine innergemeinschaftliche sonstige Leistung (B2B) gemäß Art. 196 MwSt-Richtlinie. Es wird keine österreichische Umsatzsteuer in Rechnung gestellt. Der Leistungsempfänger schuldet die Steuer im eigenen Land (Reverse Charge).', italic=True, size=9, after=Pt(8))

# --- 4. ZAHLUNGSBEDINGUNGEN ---
add_p('4. Zahlungsbedingungen', bold=True, size=13, after=Pt(6), before=Pt(8))
add_p('50 % bei Auftragserteilung (EUR 1.000,–), 50 % nach Fertigstellung & Go-Live (EUR 1.000,–).', after=Pt(2))
add_p('Zahlungsziel: 14 Tage netto.', after=Pt(8))

# --- 5. HAFTUNG ---
add_p('5. Haftungsausschluss', bold=True, size=13, after=Pt(6), before=Pt(8))
p = doc.add_paragraph()
run = p.add_run('Der Auftragnehmer haftet nicht für technische Ausfälle der Hosting- und Zahlungsanbieter (Cloudflare, Supabase, Stripe), DNS-Propagationsverzögerungen sowie für Inhalte, Bildrechte und Preisangaben, die vom Kunden bereitgestellt wurden. Indirekte Schäden und Folgeschäden sind ausgeschlossen. Der Auftragnehmer haftet ausschließlich bei vorsätzlichem oder grob fahrlässigem Verhalten, beschränkt auf den Auftragswert (EUR 2.000,–).')
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(8)

# --- 6. EIGENTUMSÜBERGANG ---
add_p('6. Eigentumsübergang', bold=True, size=13, after=Pt(6), before=Pt(8))
p = doc.add_paragraph()
run = p.add_run('Nach vollständiger Bezahlung gehen der vollständige Quellcode sowie alle erstellten Designelemente in das Eigentum des Kunden über. Der Auftragnehmer behält das Recht, das Projekt als Referenz zu nennen, sofern der Kunde nicht widerspricht.')
run.font.size = Pt(10)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(12)

# SIGNATURE
add_p('Mit freundlichen Grüßen,', after=Pt(20))
add_p('Axel Ludwig Mühlbacher', bold=True, size=11, after=Pt(2))
add_p('Geschäftsführer – Solum Invest GmbH', after=Pt(16))

# ACCEPTANCE BLOCK
add_p('Annahme des Angebots', bold=True, size=11, after=Pt(8))
add_p('Ort, Datum: ____________________________', after=Pt(4))
add_p('Ihre UID-Nr. (IT): ____________________________', after=Pt(4))
add_p('Unterschrift: ____________________________', after=Pt(2))
add_p('David Prato / Time to Jump Dolomites', after=Pt(12))

# BANK
add_p('Bankverbindung', bold=True, after=Pt(4))
add_p('Solum Invest GmbH – Raiffeisenbank Volders', size=9, after=Pt(1))
add_p('IBAN: AT08 3620 0000 0035 2971  |  BIC: RZTIAT 22200', size=9, after=Pt(1))
add_p('Verwendungszweck: Angebot 03/2026', size=9, after=Pt(8))

# FOOTER
add_p('Solum Invest GmbH · Augasse 4 · 6111 Volders · ATU81665404 · FN 642566a · office@soluminvest.at · +43 676 4008771', size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color=(128,128,128))

output_path = os.path.join(os.path.dirname(__file__), 'Angebot_03_2026_TimeToJump_Dolomites.docx')
doc.save(output_path)
print(f'Angebot saved to: {output_path}')
